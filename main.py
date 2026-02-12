from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from docx import Document
from pydantic_ai import Agent
from pydantic_ai.models.openai import OpenAIModel
import uvicorn
import tempfile
import os
from dotenv import load_dotenv

load_dotenv()

if not os.getenv("OPENAI_API_KEY"):
    raise RuntimeError("OPENAI_API_KEY is not set")

MAX_FILE_SIZE_MB = 20

app = FastAPI(title="DOCX Translator")

model = OpenAIModel("gpt-4o-mini")

translation_agent = Agent(
    model=model,
    system_prompt="""
    ترجم النص إلى العربية الفصحى.
    لا تترجم أسماء المكتبات أو المصطلحات التقنية الشائعة.
    لا تترجم أي كود أو نص داخل backticks (`).
    حافظ على تنسيق الفقرات قدر الإمكان.
    إذا كان النص كودًا فقط، أعده كما هو.
    """
)

if not os.path.exists("static"):
    os.makedirs("static")

app.mount("/static", StaticFiles(directory="static"), name="static")


@app.get("/", response_class=HTMLResponse)
async def home():
    with open("static/index.html", "r", encoding="utf-8") as f:
        return f.read()


@app.post("/process_docx")
async def process_docx(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are allowed")

    contents = await file.read()
    size_mb = len(contents) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise HTTPException(
            status_code=400,
            detail=f"File too large (max {MAX_FILE_SIZE_MB} MB)"
        )

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_in:
        tmp_in.write(contents)
        input_path = tmp_in.name

    doc = Document(input_path)
    out_doc = Document()

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            out_doc.add_paragraph("")
            continue

        result = await translation_agent.run(text)
        out_doc.add_paragraph(result.data)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_out:
        output_path = tmp_out.name
        out_doc.save(output_path)

    background_tasks.add_task(os.remove, input_path)

    return {
        "download_url": f"/download?path={output_path}"
    }


@app.get("/download")
async def download_file(
    background_tasks: BackgroundTasks,
    path: str
):
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="File not found")

    background_tasks.add_task(os.remove, path)

    return FileResponse(
        path,
        filename="translated.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
